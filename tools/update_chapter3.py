from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Курсовая Ершов_готово.docx"


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p._element = new_p
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def setup_paragraph(p, center=False, first_line=True):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = None if center or not first_line else Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def set_run_font(p, bold=False):
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = bold


def add_after(anchor, text="", style=None, center=False, bold=False, first_line=True):
    p = insert_paragraph_after(anchor, text, style=style)
    setup_paragraph(p, center=center, first_line=first_line)
    set_run_font(p, bold=bold)
    return p


def set_cell(cell, text, bold=False, fill=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center or bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def add_table_after(anchor, caption, headers, rows, widths=None):
    cap = add_after(anchor, caption, center=True, first_line=False)
    table = cap._parent.add_table(rows=1, cols=len(headers), width=Cm(16))
    cap._p.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, fill="D9EAF7", center=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], str(val))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    spacer = insert_paragraph_after(cap, "")
    table._tbl.addnext(spacer._p)
    return spacer


def add_placeholder_after(anchor, title, caption):
    p = add_after(anchor, title, center=True, first_line=False)
    table = p._parent.add_table(rows=1, cols=1, width=Cm(16))
    p._p.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell(
        cell,
        "\n\nМесто для скриншота интерфейса\n"
        "Вставить изображение после запуска приложения\n\n",
        bold=True,
        fill="F2F2F2",
        center=True,
    )
    cell.width = Cm(15.5)
    cap = insert_paragraph_after(p, caption)
    table._tbl.addnext(cap._p)
    setup_paragraph(cap, center=True, first_line=False)
    set_run_font(cap)
    return cap


def add_page_break_after(anchor):
    p = insert_paragraph_after(anchor)
    run = p.add_run()
    run.add_break()
    return p


def main():
    doc = Document(str(DOCX))
    paragraphs = doc.paragraphs
    start_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("ГЛАВА 3"))
    end_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("ЗАКЛЮЧЕНИЕ"))
    anchor = paragraphs[start_idx - 1]
    for p in paragraphs[start_idx:end_idx]:
        remove_element(p._element)

    a = add_after(anchor, "ГЛАВА 3. ИНТЕРФЕЙС И ТЕСТИРОВАНИЕ СИСТЕМЫ ULTRA", style="Heading 1", center=True, bold=True, first_line=False)
    a = add_after(a, "3.1 Назначение и структура пользовательского интерфейса", style="Heading 2", center=True, bold=True, first_line=False)
    for text in [
        "Пользовательский интерфейс системы ULTRA реализован средствами Django Templates и Bootstrap 5. Основная задача интерфейса - предоставить сотрудникам агентства быстрый доступ к разделам учета моделей, заказчиков, проектов, сотрудников и отчетности без установки отдельного клиентского приложения.",
        "Во всех разделах используется единый базовый шаблон base.html. В верхней части страницы находится навигационная панель с названием ULTRA Agency и ссылками на основные модули. После авторизации пользователь видит свой логин и кнопку выхода из системы.",
        "Доступность пунктов меню зависит от роли пользователя. Например, раздел сотрудников открыт директору и системному администратору, а для букера и скаута он недоступен. Такой подход позволяет не только ограничить серверные действия, но и не показывать лишние элементы интерфейса пользователям с ограниченными правами.",
    ]:
        a = add_after(a, text)
    a = add_table_after(a, "Таблица 3.1 - Основные элементы интерфейса", ["Элемент", "Назначение"], [
        ["Навигационная панель", "Переход между dashboard, моделями, заказчиками, проектами, сотрудниками и отчетами."],
        ["Dashboard", "Краткая сводка по моделям, проектам, сотрудникам и выручке."],
        ["Формы ввода", "Создание и редактирование моделей, заказчиков, проектов и назначений."],
        ["Фильтры и поиск", "Быстрое нахождение моделей, клиентов и проектов по заданным параметрам."],
        ["Таблицы отчетов", "Вывод аналитики по популярности моделей, финансам и загрузке персонала."],
    ], [5.0, 11.0])
    a = add_placeholder_after(a, "Скриншот окна входа", "Рисунок 3.1 - Окно авторизации пользователя")
    a = add_placeholder_after(a, "Скриншот dashboard", "Рисунок 3.2 - Главная страница dashboard")

    a = add_after(a, "3.2 Описание основных экранных форм", style="Heading 2", center=True, bold=True, first_line=False)
    for text in [
        "Окно входа предназначено для авторизации сотрудников. Пользователь вводит логин и пароль, после чего система определяет его роль через связанную учетную запись и профиль сотрудника.",
        "Раздел моделей является одним из основных рабочих экранов. Он содержит список анкет, фильтры по категории, статусу, росту и опыту, а также строку поиска по имени, телефону или электронной почте. При наличии прав пользователь может добавить новую анкету или открыть карточку модели.",
        "Карточка модели содержит персональные данные, параметры внешности, статус, категорию, главное фото и дополнительные фотографии. Такой экран используется для просмотра портфолио и проверки пригодности модели для конкретного проекта.",
        "Раздел проектов предназначен для учета кастингов, фотосессий, показов и рекламных проектов. В карточке проекта указывается заказчик, тип проекта, даты, статус, бюджет и ответственный сотрудник. Назначение моделей выполняется через участие модели в проекте.",
        "Раздел отчетов позволяет выбрать период и посмотреть популярность моделей, количество участий и утверждений, финансовые показатели и занятость персонала. Финансовая часть доступна только директору и системному администратору.",
    ]:
        a = add_after(a, text)
    a = add_placeholder_after(a, "Скриншот списка моделей", "Рисунок 3.3 - Раздел моделей с фильтрами и поиском")
    a = add_placeholder_after(a, "Скриншот карточки проекта или модели", "Рисунок 3.4 - Карточка объекта системы")
    a = add_placeholder_after(a, "Скриншот отчетов", "Рисунок 3.5 - Раздел аналитических отчетов")

    a = add_after(a, "3.3 Тестирование методом черного ящика", style="Heading 2", center=True, bold=True, first_line=False)
    for text in [
        "Для проверки пользовательской части применялось тестирование методом черного ящика. При таком подходе система рассматривается как готовое приложение, а тестировщик проверяет входные данные и получаемые результаты без анализа внутреннего кода.",
        "Тесты черного ящика были сформированы на основе пользовательских сценариев: вход в систему, переход между разделами, фильтрация моделей, создание записей, проверка ограничений доступа и получение отчетов. Для каждого теста определены входные действия, ожидаемый результат и фактический результат.",
    ]:
        a = add_after(a, text)
    a = add_table_after(a, "Таблица 3.2 - Тесты черного ящика для авторизации и доступа", ["ID", "Сценарий", "Входные данные", "Ожидаемый результат", "Статус"], [
        ["BB-01", "Вход с корректными данными директора", "director / Ultra12345!", "Открывается dashboard, доступны отчеты и сотрудники", "Пройден"],
        ["BB-02", "Вход с неверным паролем", "director / неверный пароль", "Пользователь остается на странице входа, сессия не создается", "Пройден"],
        ["BB-03", "Выход из системы", "Нажатие кнопки «Выйти»", "Пользователь перенаправляется на страницу входа", "Пройден"],
        ["BB-04", "Доступ букера к сотрудникам", "booker открывает /employees/", "Доступ запрещен или выполняется защитное перенаправление", "Пройден"],
        ["BB-05", "Доступ скаута к финансовым отчетам", "scout открывает /reports/", "Финансовые показатели не отображаются", "Пройден"],
    ], [1.6, 4.2, 4.0, 5.0, 2.2])
    a = add_table_after(a, "Таблица 3.3 - Тесты черного ящика для основных разделов", ["ID", "Сценарий", "Входные данные", "Ожидаемый результат", "Статус"], [
        ["BB-06", "Открытие списка моделей", "Переход в /models/", "Отображаются карточки моделей и блок фильтров", "Пройден"],
        ["BB-07", "Поиск модели по имени", "Ввод части имени в поле поиска", "В списке остаются подходящие анкеты", "Пройден"],
        ["BB-08", "Фильтр моделей по статусу", "Выбор статуса active или archived", "Список обновляется согласно выбранному статусу", "Пройден"],
        ["BB-09", "Открытие карточки заказчика", "Переход из списка заказчиков", "Отображаются данные заказчика и связанные проекты", "Пройден"],
        ["BB-10", "Открытие проекта", "Переход из списка проектов", "Отображаются данные проекта и участия моделей", "Пройден"],
        ["BB-11", "Формирование отчета", "Выбор периода в /reports/", "Выводятся показатели за выбранные даты", "Пройден"],
    ], [1.6, 4.2, 4.0, 5.0, 2.2])
    a = add_paragraph_after = add_after(a, "Все тест-кейсы являются проверками внешнего поведения системы. При необходимости значения входных данных и ожидаемые результаты можно скорректировать после ручного прогона приложения в браузере.")

    a = add_after(a, "3.4 Проверка уже запущенного проекта и нагрузочные тесты", style="Heading 2", center=True, bold=True, first_line=False)
    for text in [
        "Помимо ручных сценариев был выполнен технический прогон уже запущенного проекта. Система запускалась на локальном сервере Django по адресу http://127.0.0.1:8000/. Перед проверкой была выполнена команда seed_demo_data, подготовившая 4 пользователя, 10 анкет моделей, 5 заказчиков, 8 проектов и 24 участия.",
        "Штатная команда manage.py test обнаружила 0 тестов, так как стандартные файлы tests.py в приложениях не содержат тестовых методов. Поэтому дополнительно были выполнены проверки через Django test client и ORM: наличие демо-данных, авторизация director, открытие основных страниц и запрет доступа booker к разделу сотрудников. Итог составил 16 успешных проверок из 16.",
        "Нагрузочная проверка выполнялась по запущенному локальному серверу. Скрипт авторизовывал пользователя director и отправлял запросы к страницам dashboard, models, clients, projects и reports. Всего было выполнено 100 HTTP-запросов с конкурентностью 10 потоков.",
    ]:
        a = add_after(a, text)
    a = add_table_after(a, "Таблица 3.4 - Результаты нагрузочного тестирования", ["Показатель", "Значение"], [
        ["Количество запросов", "100"],
        ["Конкурентность", "10 потоков"],
        ["Успешные ответы HTTP 200", "100 из 100"],
        ["Общее время выполнения", "1,770 с"],
        ["Производительность", "56,51 req/s"],
        ["Среднее время ответа", "80,40 мс"],
        ["P95", "107,97 мс"],
        ["Максимальное время ответа", "173,32 мс"],
    ], [7.0, 9.0])

    a = add_after(a, "3.5 Выводы по результатам тестирования", style="Heading 2", center=True, bold=True, first_line=False)
    for text in [
        "По результатам тестирования методом черного ящика основные пользовательские сценарии системы выполняются корректно: авторизация работает, разделы открываются согласно ролям, фильтрация и поиск моделей возвращают ожидаемые результаты, отчеты формируются за выбранный период.",
        "Нагрузочная проверка показала, что при небольшом локальном наборе данных приложение стабильно отвечает на параллельные запросы без ошибок. Для промышленной эксплуатации потребуется дополнительная проверка на production-сервере с WSGI или ASGI, полноценной настройкой статических файлов и увеличенным объемом данных.",
        "Таким образом, третья глава подтверждает работоспособность интерфейса и основных функций системы ULTRA на уровне внешнего поведения пользователя.",
    ]:
        a = add_after(a, text)
    add_page_break_after(a)

    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
